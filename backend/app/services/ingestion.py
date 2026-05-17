import os
import shutil
import fitz  # PyMuPDF
import docx
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from app.models.document import Document, DocumentChunk
from app.schemas.document import DocumentCreate
from app.crud import crud_document
import logging

logger = logging.getLogger(__name__)

UPLOAD_DIR = "uploads"
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
ALLOWED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "md"
}

os.makedirs(UPLOAD_DIR, exist_ok=True)

def extract_pages_from_file(file_path: str, mime_type: str) -> list:
    pages = []
    try:
        if mime_type == "application/pdf":
            doc = fitz.open(file_path)
            for i, page in enumerate(doc):
                pages.append({"text": page.get_text(), "page_number": i + 1})
            doc.close()
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            pages.append({"text": text, "page_number": 1})
        elif mime_type in ["text/plain", "text/markdown"]:
            with open(file_path, "r", encoding="utf-8") as f:
                pages.append({"text": f.read(), "page_number": 1})
        else:
            raise ValueError("Unsupported file type for extraction")
        return pages
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        raise e

from app.services.chunking import chunk_text

def process_upload(file: UploadFile, user_id: int, db: Session) -> Document:
    # 1. Validation
    if file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported file type.")
    
    file.file.seek(0, 2)
    file_size = file.file.tell()
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Max size is 10MB.")
    file.file.seek(0)

    # 2. Save File
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        logger.error(f"Failed to save file: {e}")
        raise HTTPException(status_code=500, detail="Could not save file.")

    # 3. Create initial DB record
    db_doc = Document(
        title=file.filename,
        filename=file.filename,
        file_type=file.content_type,
        access_level="private",
        uploaded_by=user_id,
        status="uploaded"
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)

    # 4. Extract Text & Chunk
    try:
        db_doc.status = "processing"
        db.commit()

        pages = extract_pages_from_file(file_path, db_doc.file_type)
        
        chunk_index = 0
        for page in pages:
            chunks_data = chunk_text(
                text=page["text"],
                document_id=db_doc.id,
                filename=db_doc.filename,
                page_number=page.get("page_number"),
                chunk_index=chunk_index,
                start_chunk_index=chunk_index
            )
            
            for chunk_data in chunks_data:
                db_chunk = DocumentChunk(
                    document_id=db_doc.id,
                    chunk_index=chunk_data["chunk_index"],
                    content=chunk_data["content"],
                    metadata_={
                        "filename": chunk_data["filename"],
                        "page_number": chunk_data["page_number"],
                        "section_title": chunk_data.get("section_title")
                    }
                )
                db.add(db_chunk)
            
            chunk_index += len(chunks_data)
        
        db_doc.status = "indexed" 
        db.commit()
        db.refresh(db_doc)

    except Exception as e:
        logger.error(f"Processing failed for document {db_doc.id}: {e}")
        db_doc.status = "failed"
        db.commit()
        raise HTTPException(status_code=500, detail="Failed to process document text.")

    return db_doc
