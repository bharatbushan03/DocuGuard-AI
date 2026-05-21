import os
import shutil
import fitz  # PyMuPDF
import docx
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from app.models.document import Document, DocumentChunk
from app.db.session import SessionLocal
from app.services.chunking import chunk_text
from app.services.embedding import embed_batch
from app.services.vector_db import store_vectors
from app.core.file_validation import ALLOWED_MIME_TO_EXT, validate_upload_content
import logging
import uuid

logger = logging.getLogger(__name__)

UPLOAD_DIR = "uploads"
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
ALLOWED_MIME_TYPES = ALLOWED_MIME_TO_EXT
MAX_PDF_PAGES = 500

os.makedirs(UPLOAD_DIR, exist_ok=True)

def extract_pages_from_file(file_path: str, mime_type: str) -> list:
    pages = []
    try:
        if mime_type == "application/pdf":
            doc = fitz.open(file_path)
            if doc.page_count > MAX_PDF_PAGES:
                doc.close()
                raise ValueError(f"PDF exceeds maximum page limit ({MAX_PDF_PAGES}).")
            for i, page in enumerate(doc):
                pages.append({"text": page.get_text(), "page_number": i + 1})
            doc.close()
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            pages.append({"text": text, "page_number": 1})
        elif mime_type in ["text/plain", "text/markdown"]:
            with open(file_path, "r", encoding="utf-8") as f:
                pages.append({"text": f.read(), "page_number": 1})
        else:
            raise ValueError("Unsupported file type for extraction")
        return pages
    except Exception as e:
        logger.error("Error extracting text from document_id path=%s: %s", file_path, e)
        raise e

def save_uploaded_file(file: UploadFile, user_id: int, db: Session) -> Document:
    declared_mime = file.content_type or ""
    if declared_mime not in ALLOWED_MIME_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported file type.")

    file.file.seek(0, 2)
    file_size = file.file.tell()
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Max size is 10MB.")
    file.file.seek(0)

    content = file.file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Max size is 10MB.")

    try:
        verified_mime, file_extension = validate_upload_content(content, declared_mime)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    safe_filename = f"{uuid.uuid4()}.{file_extension}"
    file_path = os.path.join(UPLOAD_DIR, safe_filename)
    try:
        with open(file_path, "wb") as buffer:
            buffer.write(content)
    except Exception as e:
        logger.error("Failed to save upload for user_id=%s: %s", user_id, e)
        raise HTTPException(status_code=500, detail="Could not save file.")

    # 3. Create initial DB record
    db_doc = Document(
        title=file.filename, # Keep original filename as title
        filename=safe_filename, # Safe filename for storage
        file_type=verified_mime,
        access_level="private",
        uploaded_by=user_id,
        status="uploaded"
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)

    return db_doc, file_path

def process_document_background(document_id: int, file_path: str):
    """
    Background task to extract text, chunk, embed, and store in Qdrant.
    """
    logger.info(f"Starting background processing for document {document_id}")
    db = SessionLocal()
    try:
        db_doc = db.query(Document).filter(Document.id == document_id).first()
        if not db_doc:
            logger.error(f"Document {document_id} not found.")
            return

        db_doc.status = "processing"
        db.commit()

        # Extract Text
        pages = extract_pages_from_file(file_path, db_doc.file_type)
        
        # Chunk Text
        chunk_index = 0
        all_chunks_data = []
        for page in pages:
            chunks_data = chunk_text(
                text=page["text"],
                document_id=db_doc.id,
                filename=db_doc.title,
                page_number=page.get("page_number"),
                start_chunk_index=chunk_index
            )
            all_chunks_data.extend(chunks_data)
            chunk_index += len(chunks_data)
            
        # Save chunks to DB
        db_chunks = []
        for chunk_data in all_chunks_data:
            db_chunk = DocumentChunk(
                document_id=db_doc.id,
                chunk_index=chunk_data["chunk_index"],
                content=chunk_data["content"],
                metadata_={
                    "filename": chunk_data["filename"],
                    "page_number": chunk_data["page_number"],
                    "section_title": chunk_data.get("section_title")
                }
            )
            db.add(db_chunk)
            db_chunks.append(db_chunk)
            
        db.commit()
        for chunk in db_chunks:
            db.refresh(chunk)
            
        # Embed and store in Qdrant
        for idx, chunk_data in enumerate(all_chunks_data):
            chunk_data["chunk_id"] = db_chunks[idx].id
            chunk_data["access_level"] = db_doc.access_level
            chunk_data["uploaded_by"] = db_doc.uploaded_by
            
        texts_to_embed = [c["content"] for c in all_chunks_data]
        
        batch_size = 100
        all_embeddings = []
        for i in range(0, len(texts_to_embed), batch_size):
            batch_texts = texts_to_embed[i:i + batch_size]
            embeddings = embed_batch(batch_texts)
            all_embeddings.extend(embeddings)
            
        point_ids = store_vectors(all_chunks_data, all_embeddings)
        
        # Update DB with Qdrant Point IDs
        for idx, point_id in enumerate(point_ids):
            db_chunks[idx].qdrant_point_id = point_id
            
        db_doc.status = "indexed" 
        db.commit()
        logger.info(f"Successfully processed document {document_id}")

    except Exception as e:
        logger.error("Processing failed for document_id=%s: %s", document_id, e)
        db_doc = db.query(Document).filter(Document.id == document_id).first()
        if db_doc:
            db_doc.status = "failed"
            db.commit()
    finally:
        db.close()
